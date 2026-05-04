from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\TechMarket")
OUT = ROOT / "TechMarket_Server_Foundation_Notes.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return paragraph


def add_note(document, title, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E6FFFA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(15, 118, 110)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
    document.add_paragraph()


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "0F766E")
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=(255, 255, 255))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph()
    return table


def style_document(document):
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for name, size, color in [
        ("Heading 1", 18, (15, 118, 110)),
        ("Heading 2", 14, (23, 32, 51)),
        ("Heading 3", 12, (23, 32, 51)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def main():
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TechMarket Server Foundation")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Памятка по первому этапу разработки backend на NestJS")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)

    add_note(
        document,
        "Как использовать документ",
        "Этот файл можно читать как учебную памятку и использовать как черновую основу для диплома. Он объясняет не только что было добавлено, но и зачем это нужно в архитектуре интернет-магазина.",
    )

    document.add_heading("1. Что уже сделано", level=1)
    add_bullets(
        document,
        [
            "Сервер запускается как REST API на NestJS.",
            "Все маршруты имеют общий префикс /api.",
            "Подключена конфигурация через .env.",
            "Подключен PostgreSQL через Prisma ORM.",
            "Включена глобальная валидация входных данных.",
            "Подключена Swagger/OpenAPI документация.",
            "Добавлен endpoint проверки работоспособности сервера.",
        ],
    )

    document.add_heading("2. Структура проекта", level=1)
    add_table(
        document,
        ["Часть", "Путь", "Назначение"],
        [
            ["Server", "apps/server", "Backend на NestJS"],
            ["Prisma", "prisma", "Схема базы данных и будущие миграции"],
            ["Env", ".env", "Настройки окружения"],
            ["Client", "apps/client", "Будущий frontend на React + Vite"],
        ],
    )
    document.add_paragraph(
        "Главная идея архитектуры: frontend и backend развиваются отдельно. Клиент отправляет HTTP-запросы на сервер, сервер работает с базой данных и возвращает JSON."
    )

    document.add_heading("3. Как проходит запрос", level=1)
    document.add_paragraph("Пример запроса проверки сервера:")
    add_code(document, "Браузер -> GET http://localhost:5000/api/health\nNestJS -> AppController\nAppController -> AppService\nAppService -> JSON response")
    document.add_paragraph("Ответ сервера:")
    add_code(document, '{\n  "status": "ok",\n  "service": "TechMarket API"\n}')

    document.add_heading("4. main.ts - точка входа", level=1)
    document.add_paragraph("Файл apps/server/src/main.ts отвечает за запуск NestJS-приложения.")
    add_bullets(
        document,
        [
            "создает приложение через NestFactory.create(AppModule);",
            "берет порт сервера из .env;",
            "задает глобальный префикс /api;",
            "включает CORS для будущего frontend;",
            "подключает глобальную валидацию DTO;",
            "подключает Swagger-документацию.",
        ],
    )
    add_note(
        document,
        "Формулировка для диплома",
        "Точка входа серверного приложения выполняет инициализацию NestJS, задает общий префикс API, включает CORS для взаимодействия с клиентской частью, подключает глобальную валидацию входных данных и формирует OpenAPI-документацию.",
    )

    document.add_heading("5. Глобальный префикс /api", level=1)
    document.add_paragraph("В коде используется:")
    add_code(document, 'app.setGlobalPrefix("api");')
    document.add_paragraph("Это значит, что все endpoints начинаются с /api:")
    add_code(document, "health -> /api/health\nproducts -> /api/products\nauth/login -> /api/auth/login")

    document.add_heading("6. CORS", level=1)
    document.add_paragraph("CORS нужен, чтобы frontend мог обращаться к backend.")
    add_table(
        document,
        ["Адрес", "Зачем нужен"],
        [
            ["http://localhost:5173", "Стандартный адрес Vite frontend"],
            ["http://localhost:3000", "Совместимость с предыдущим прототипом"],
        ],
    )
    document.add_paragraph("Параметр credentials: true понадобится для авторизации через HTTP-only cookies.")

    document.add_heading("7. ValidationPipe", level=1)
    add_table(
        document,
        ["Настройка", "Что делает"],
        [
            ["whitelist: true", "Удаляет лишние поля из запроса"],
            ["forbidNonWhitelisted: true", "Запрещает поля, которых нет в DTO"],
            ["transform: true", "Преобразует значения к нужным типам"],
        ],
    )
    document.add_paragraph(
        "Это понадобится для DTO регистрации, создания товара, создания категории и оформления заказа. Такая проверка повышает безопасность и снижает риск ошибок."
    )

    document.add_heading("8. Swagger/OpenAPI", level=1)
    document.add_paragraph("Swagger доступен по адресу:")
    add_code(document, "http://localhost:5000/api/docs")
    document.add_paragraph(
        "Он нужен для просмотра и тестирования API. В дипломе это удобно показывать как интерактивную документацию серверной части."
    )

    document.add_heading("9. AppModule", level=1)
    document.add_paragraph("Файл apps/server/src/app.module.ts - корневой модуль приложения. Сейчас он подключает:")
    add_bullets(document, ["ConfigModule", "PrismaModule", "AppController", "AppService"])
    document.add_paragraph("В будущем сюда будут добавляться AuthModule, UsersModule, CategoriesModule, ProductsModule, OrdersModule и другие модули.")

    document.add_heading("10. ConfigModule и .env", level=1)
    add_table(
        document,
        ["Переменная", "Назначение"],
        [
            ["PORT", "Порт сервера"],
            ["DATABASE_URL", "Строка подключения к PostgreSQL"],
            ["JWT_ACCESS_SECRET", "Секрет для access token"],
            ["JWT_REFRESH_SECRET", "Секрет для refresh token"],
        ],
    )
    add_note(
        document,
        "Важно",
        "Пароль базы данных не должен попадать в публичный репозиторий. В учебном проекте он может лежать локально в .env, но в реальном проекте используется защищенное окружение сервера.",
    )

    document.add_heading("11. Prisma ORM", level=1)
    document.add_paragraph(
        "Prisma - это ORM, то есть слой между TypeScript-кодом и базой данных. Вместо ручного SQL сервер будет обращаться к базе через Prisma Client."
    )
    add_code(document, "this.prisma.product.findMany()")
    document.add_paragraph("Это удобнее, типобезопаснее и снижает риск SQL-инъекций.")

    document.add_heading("12. Prisma schema", level=1)
    document.add_paragraph("Файл prisma/schema.prisma сейчас содержит базовую настройку:")
    add_code(document, 'generator client {\n  provider = "prisma-client-js"\n}\n\ndatasource db {\n  provider = "postgresql"\n}')
    document.add_paragraph("Модели базы данных будут добавляться на следующих этапах.")

    document.add_heading("13. PrismaService", level=1)
    document.add_paragraph("Файлы:")
    add_code(document, "apps/server/src/prisma/prisma.module.ts\napps/server/src/prisma/prisma.service.ts")
    document.add_paragraph("PrismaService делает Prisma доступной внутри NestJS. В будущих сервисах можно будет писать:")
    add_code(document, "constructor(private prisma: PrismaService) {}")

    document.add_heading("14. Особенность Prisma 7", level=1)
    document.add_paragraph("В проекте используется Prisma 7, поэтому для PostgreSQL нужен adapter:")
    add_code(document, "@prisma/adapter-pg")
    document.add_paragraph("PrismaService получает DATABASE_URL из ConfigService и передает его в adapter:")
    add_code(document, 'const adapter = new PrismaPg({\n  connectionString: configService.getOrThrow<string>("DATABASE_URL"),\n});\n\nsuper({ adapter });')

    document.add_heading("15. Health endpoint", level=1)
    document.add_paragraph("Health endpoint нужен для быстрой проверки, что сервер работает.")
    add_code(document, "GET /api/health")
    document.add_paragraph("Если endpoint отвечает, значит сервер стартовал и базовая инфраструктура работает.")

    document.add_heading("16. Как запускать сервер", level=1)
    add_code(document, "cd C:\\TechMarket\\apps\\server\nnpm.cmd run start:dev")
    document.add_paragraph("Проверить в браузере:")
    add_code(document, "http://localhost:5000/api/health\nhttp://localhost:5000/api/docs")

    document.add_heading("17. Полезные команды", level=1)
    add_table(
        document,
        ["Команда", "Где выполнять", "Назначение"],
        [
            ["npx.cmd prisma generate", "C:\\TechMarket", "Сгенерировать Prisma Client"],
            ["npx.cmd prisma validate", "C:\\TechMarket", "Проверить Prisma-схему"],
            ["npm.cmd run build", "apps/server", "Проверить сборку сервера"],
            ["npm.cmd run start:dev", "apps/server", "Запустить сервер в dev-режиме"],
        ],
    )

    document.add_heading("18. Текст для диплома", level=1)
    document.add_paragraph(
        "Серверная часть системы реализована на NestJS с использованием модульной архитектуры. На начальном этапе был настроен базовый инфраструктурный слой: глобальный префикс REST API, CORS, валидация входных данных, загрузка переменных окружения, подключение Prisma ORM к PostgreSQL и автоматическая документация Swagger/OpenAPI. Такой подход обеспечивает расширяемость серверной части и упрощает дальнейшую реализацию модулей авторизации, каталога, корзины, заказов и административной панели."
    )

    document.add_heading("19. Следующий шаг", level=1)
    document.add_paragraph("Рекомендуемый следующий учебный шаг - модуль категорий.")
    add_bullets(
        document,
        [
            "он проще, чем авторизация;",
            "показывает схему Controller -> Service -> Prisma;",
            "нужен для товаров и шаблонов характеристик;",
            "удобно тестируется через Swagger.",
        ],
    )
    add_code(document, "GET /api/categories\nGET /api/categories/:id\nPOST /api/categories\nPATCH /api/categories/:id\nDELETE /api/categories/:id")

    document.add_heading("20. Мини-глоссарий", level=1)
    add_table(
        document,
        ["Термин", "Простое объяснение"],
        [
            ["NestJS", "Backend-фреймворк для Node.js"],
            ["Module", "Часть приложения, объединяющая controller и service"],
            ["Controller", "Принимает HTTP-запросы"],
            ["Service", "Содержит бизнес-логику"],
            ["DTO", "Описание данных, которые приходят в запросе"],
            ["ValidationPipe", "Автоматическая проверка DTO"],
            ["Prisma", "ORM для работы с базой данных"],
            ["PostgreSQL", "База данных"],
            ["Swagger", "Интерактивная документация API"],
            ["CORS", "Разрешение frontend обращаться к backend"],
            [".env", "Файл с настройками окружения"],
        ],
    )

    document.add_heading("21. Краткая схема", level=1)
    add_code(document, "Frontend React + Vite\n        |\n        | HTTP request\n        v\nNestJS Controller\n        |\n        v\nNestJS Service\n        |\n        v\nPrismaService\n        |\n        v\nPostgreSQL")

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
